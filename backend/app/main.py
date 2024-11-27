from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from pydantic import BaseModel
from datetime import datetime
import os
from typing import Optional, List, Dict
import io
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.core.database import Base, engine, SessionLocal, get_db
from app.core.models import Document

# Pydantic models for API
class DocumentResponse(BaseModel):
    id: int
    title: str
    content: Optional[str]
    file_type: str
    created_at: datetime

    class Config:
        from_attributes = True

class DocumentList(BaseModel):
    total: int
    documents: List[DocumentResponse]

Base.metadata.create_all(bind=engine)

app = FastAPI(title="Knowledge Window Backend")

# Initialize services
from .services.vector_search import VectorSearchService
from .services.notion_service import NotionService
from .services.confluence_service import ConfluenceService

vector_search = VectorSearchService()
notion_service = NotionService(os.getenv("NOTION_TOKEN", ""))
confluence_service = ConfluenceService(
    os.getenv("CONFLUENCE_URL", ""),
    os.getenv("CONFLUENCE_USER", ""),
    os.getenv("CONFLUENCE_TOKEN", "")
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Cache with TTL
document_cache = {}
CACHE_TTL = 3600  # 1 hour

# Document processing
def extract_text_from_pdf(file_content: bytes) -> str:
    pdf = PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_content))
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

@app.post("/api/v1/documents", response_model=DocumentResponse)
async def create_document(title: str, file: UploadFile = File(...), db: Session = Depends(get_db)):
    content = await file.read()
    file_ext = file.filename.split('.')[-1].lower()

    if file_ext not in ['pdf', 'docx', 'txt']:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    # Process different file types
    if file_ext == 'pdf':
        content_text = extract_text_from_pdf(content)
    elif file_ext == 'docx':
        content_text = extract_text_from_docx(content)
    else:
        content_text = content.decode()

    doc = Document(
        title=title,
        content=content_text,
        file_type=file_ext
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return doc

@app.get("/api/v1/documents/{doc_id}", response_model=DocumentResponse)
def get_document(doc_id: int, db: Session = Depends(get_db)):
    if doc_id in document_cache:
        return document_cache[doc_id]

    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    document_cache[doc_id] = doc
    return doc

@app.get("/api/v1/documents", response_model=DocumentList)
def get_documents(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    cache_key = f"docs_{skip}_{limit}"
    if cache_key in document_cache:
        return document_cache[cache_key]

    total = db.query(Document).count()
    docs = db.query(Document).offset(skip).limit(limit).all()

    result = DocumentList(
        total=total,
        documents=docs
    )
    document_cache[cache_key] = result
    return result

@app.get("/api/v1/search", response_model=List[DocumentResponse])
async def search_documents(
    query: str,
    limit: int = 10,
    search_type: str = "text",  # "text" or "vector"
    db: Session = Depends(get_db)
):
    cache_key = f"search_{search_type}_{query}_{limit}"
    if cache_key in document_cache:
        return document_cache[cache_key]

    if search_type == "vector":
        # Ensure vectors are up to date
        vector_search.index_documents(db)
        doc_scores = vector_search.search(query, limit)
        doc_ids = [doc_id for doc_id, _ in doc_scores]
        docs = db.query(Document).filter(Document.id.in_(doc_ids)).all()
    else:
        docs = db.query(Document).filter(
            Document.content.ilike(f"%{query}%")
        ).limit(limit).all()

    document_cache[cache_key] = docs
    return docs

@app.post("/api/v1/import/notion")
async def import_notion_page(
    page_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    try:
        page_content = notion_service.get_page_content(page_id)
        doc = Document(
            title=f"Notion: {page_id}",
            content=page_content["content"],
            file_type="notion"
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        # Update vector index in background
        background_tasks.add_task(vector_search.index_documents, db)

        return {"id": doc.id, "status": "imported"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/v1/import/confluence")
async def import_confluence_page(
    page_id: str,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db)
):
    try:
        page_data = confluence_service.get_page_content(page_id)
        doc = Document(
            title=page_data["title"],
            content=page_data["content"],
            file_type="confluence"
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        # Update vector index in background
        background_tasks.add_task(vector_search.index_documents, db)

        return {"id": doc.id, "status": "imported"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
