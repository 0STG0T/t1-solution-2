from typing import Optional, BinaryIO
import io
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session
from ..core.models import Document

class DocumentProcessor:
    def __init__(self, db: Session):
        self.db = db

    def process_document(self, file: BinaryIO, title: str, file_type: str) -> Document:
        content = file.read()
        text_content = self._extract_text(content, file_type)

        document = Document(
            title=title,
            content=text_content,
            raw_content=content,
            file_type=file_type
        )

        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)
        return document

    def _extract_text(self, content: bytes, file_type: str) -> str:
        if file_type == 'pdf':
            return self._extract_pdf_text(content)
        elif file_type == 'docx':
            return self._extract_docx_text(content)
        elif file_type == 'txt':
            return content.decode('utf-8')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf_text(self, content: bytes) -> str:
        pdf = PdfReader(io.BytesIO(content))
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
        return text

    def _extract_docx_text(self, content: bytes) -> str:
        doc = DocxDocument(io.BytesIO(content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
